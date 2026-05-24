"""
Supplemental extraction for files marked as warning/error in extraction_report.csv.

The script is intentionally conservative:
- PDFs are re-extracted with pdfplumber, including page text and detected tables.
- DOCX image containers reuse a successful sibling extraction when available.
- Empty or temporary Word files are recorded as excluded.
"""

from __future__ import annotations

import csv
import json
import re
import sys
import zipfile
from collections import defaultdict
from pathlib import Path

from docx import Document  # type: ignore
from pypdf import PdfReader  # type: ignore


DATA_DIR = Path("data")
EXTRACTED_DIR = Path("extracted_texts")
REPORT = Path("extraction_report.csv")
SUPPLEMENTAL_REPORT = Path("supplemental_processing_report.csv")
CHECKPOINT = Path("embed_checkpoint.json")


def is_success(status: str) -> bool:
    return bool(status) and ord(status[0]) == 0x2705


def output_path_for(source_rel: str) -> Path:
    rel = Path(source_rel)
    return EXTRACTED_DIR / rel.with_suffix(".txt")


def read_existing_text_for(source_rel: str) -> str:
    out = output_path_for(source_rel)
    if out.exists():
        return out.read_text(encoding="utf-8", errors="replace")
    return ""


def write_text(source_rel: str, content: str) -> Path:
    out = output_path_for(source_rel)
    out.parent.mkdir(parents=True, exist_ok=True)
    out.write_text(content.strip() + "\n", encoding="utf-8")
    return out


def clean_cell(value: object) -> str:
    if value is None:
        return ""
    return re.sub(r"\s+", " ", str(value)).strip()


def markdown_table(table: list[list[object]]) -> str:
    rows = [[clean_cell(cell) for cell in row] for row in table if row]
    rows = [row for row in rows if any(row)]
    if not rows:
        return ""
    width = max(len(row) for row in rows)
    rows = [row + [""] * (width - len(row)) for row in rows]
    header = rows[0]
    body = rows[1:]
    lines = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join(["---"] * width) + " |",
    ]
    lines.extend("| " + " | ".join(row) + " |" for row in body)
    return "\n".join(lines)


def extract_pdf(path: Path) -> tuple[str, dict[str, int]]:
    sections: list[str] = []
    page_count = 0
    layout_pages = 0
    reader = PdfReader(str(path))
    page_count = len(reader.pages)
    for page_index, page in enumerate(reader.pages, 1):
        page_parts: list[str] = []
        text = page.extract_text() or ""
        layout_text = ""
        try:
            layout_text = page.extract_text(extraction_mode="layout") or ""
        except TypeError:
            layout_text = ""
        if text.strip():
            page_parts.append("[Text extraction]\n" + text.strip())
        if layout_text.strip() and layout_text.strip() != text.strip():
            layout_pages += 1
            page_parts.append("[Layout-preserving extraction]\n" + layout_text.strip())
        if page_parts:
            sections.append(f"--- Page {page_index} ---\n" + "\n\n".join(page_parts))
    return "\n\n".join(sections), {"pages": page_count, "layout_pages": layout_pages}


def extract_docx(path: Path) -> tuple[str, dict[str, int]]:
    doc = Document(path)
    parts: list[str] = []
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    if paragraphs:
        parts.append("\n".join(paragraphs))
    table_count = 0
    for table in doc.tables:
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        rendered = markdown_table(rows)
        if rendered:
            table_count += 1
            parts.append(f"[DOCX Table {table_count}]\n{rendered}")
    image_count = count_docx_images(path)
    if image_count:
        parts.append(f"[NOTE] DOCX contains {image_count} embedded image(s); OCR was not run here.")
    return "\n\n".join(parts), {"tables": table_count, "images": image_count}


def count_docx_images(path: Path) -> int:
    try:
        with zipfile.ZipFile(path) as zf:
            return sum(1 for name in zf.namelist() if name.startswith("word/media/"))
    except zipfile.BadZipFile:
        return 0


def sibling_success_text(row: list[str], rows_by_dir: dict[str, list[list[str]]]) -> tuple[str, str]:
    parent = str(Path(row[0]).parent)
    stem = Path(row[0]).stem.replace(" 圖片", "").replace("圖片", "")
    candidates: list[tuple[int, list[str]]] = []
    for other in rows_by_dir[parent]:
        if other is row or not is_success(other[2]):
            continue
        other_stem = Path(other[0]).stem.replace(" 圖片", "").replace("圖片", "")
        score = 0
        if other_stem == stem:
            score += 10
        if stem and (stem in other_stem or other_stem in stem):
            score += 5
        if other[1].lower() == ".txt":
            score += 3
        if other[1].lower() == ".pdf":
            score += 2
        if score:
            candidates.append((score, other))
    candidates.sort(key=lambda x: x[0], reverse=True)
    for _, candidate in candidates:
        text = read_existing_text_for(candidate[0])
        if text.strip():
            return text, candidate[0]
    return "", ""


def update_checkpoint_remove(paths: list[Path]) -> int:
    if not CHECKPOINT.exists():
        return 0
    done = set(json.loads(CHECKPOINT.read_text(encoding="utf-8")))
    before = len(done)
    for path in paths:
        done.discard(str(path))
    CHECKPOINT.write_text(
        json.dumps(sorted(done), ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    return before - len(done)


def main() -> None:
    with REPORT.open(encoding="utf-8-sig", newline="") as f:
        rows = list(csv.reader(f))

    header, data = rows[0], rows[1:]
    original_status = {}
    if SUPPLEMENTAL_REPORT.exists():
        with SUPPLEMENTAL_REPORT.open(encoding="utf-8-sig", newline="") as f:
            for old in csv.DictReader(f):
                original_status.setdefault(old["檔案路徑"], old["原狀態"])

    rows_by_dir: dict[str, list[list[str]]] = defaultdict(list)
    for row in data:
        if row[0] in original_status:
            row[2] = original_status[row[0]]
        rows_by_dir[str(Path(row[0]).parent)].append(row)

    supplemental_rows: list[dict[str, str]] = []
    rewritten_outputs: list[Path] = []

    for row in data:
        status = row[2]
        if is_success(status):
            continue

        source_rel = row[0]
        source_path = DATA_DIR / source_rel
        action = ""
        new_status = status
        detail = ""
        out_path = ""

        try:
            if source_path.name.startswith("~$"):
                action = "excluded"
                new_status = "⏭️排除-Word暫存檔"
                detail = "Word temporary lock file; not a real source document."
            elif not source_path.exists() or source_path.stat().st_size == 0:
                action = "excluded"
                new_status = "⏭️排除-空白或不存在"
                detail = "Source file is empty or missing."
            elif source_path.suffix.lower() == ".pdf":
                text, meta = extract_pdf(source_path)
                if text.strip():
                    content = (
                        f"[補處理來源] {source_rel}\n"
                        f"[補處理方法] pypdf text/layout extraction\n"
                        f"[頁數] {meta.get('pages', 0)}\n"
                        f"[版面保留頁數] {meta.get('layout_pages', 0)}\n\n"
                        f"{text}"
                    )
                    out = write_text(source_rel, content)
                    rewritten_outputs.append(out)
                    out_path = str(out)
                    if status == "⚠️表格型":
                        new_status = "✅補處理-表格抽取"
                    else:
                        new_status = "✅補處理-短公告確認"
                    action = "rewritten"
                    detail = f"Extracted {len(text)} chars, pages={meta.get('pages', 0)}, layout_pages={meta.get('layout_pages', 0)}."
                else:
                    action = "needs_manual"
                    new_status = "⚠️補處理後仍無文字"
                    detail = "pdfplumber could not extract text or tables."
            elif source_path.suffix.lower() == ".docx":
                text, sibling = sibling_success_text(row, rows_by_dir)
                if text.strip():
                    content = (
                        f"[補處理來源] {source_rel}\n"
                        f"[補處理方法] reuse successful sibling extraction\n"
                        f"[對應成功檔] {sibling}\n\n"
                        f"{text}"
                    )
                    out = write_text(source_rel, content)
                    rewritten_outputs.append(out)
                    out_path = str(out)
                    new_status = "✅補處理-使用對應成功檔"
                    action = "rewritten_from_sibling"
                    detail = f"Reused text from {sibling}."
                else:
                    text, meta = extract_docx(source_path)
                    if text.strip():
                        content = (
                            f"[補處理來源] {source_rel}\n"
                            f"[補處理方法] python-docx text/table extraction\n"
                            f"[表格數] {meta.get('tables', 0)}\n"
                            f"[圖片數] {meta.get('images', 0)}\n\n"
                            f"{text}"
                        )
                        out = write_text(source_rel, content)
                        rewritten_outputs.append(out)
                        out_path = str(out)
                        new_status = "✅補處理-DOCX抽取"
                        action = "rewritten"
                        detail = f"Extracted {len(text)} chars."
                    else:
                        action = "needs_manual"
                        new_status = "⚠️補處理後仍無文字"
                        detail = "DOCX has no extractable text and no successful sibling."
            else:
                action = "needs_manual"
                new_status = "⚠️補處理未支援格式"
                detail = f"Unsupported extension: {source_path.suffix}"
        except Exception as exc:
            action = "error"
            new_status = "❌補處理錯誤"
            detail = repr(exc)

        row[2] = new_status
        if len(row) >= 7:
            row[6] = (row[6] + "；" if row[6] else "") + f"補處理：{detail}"

        supplemental_rows.append({
            "檔案路徑": source_rel,
            "原狀態": status,
            "新狀態": new_status,
            "處理動作": action,
            "輸出檔": out_path,
            "說明": detail,
        })

    with REPORT.open("w", encoding="utf-8-sig", newline="") as f:
        csv.writer(f).writerows([header] + data)

    with SUPPLEMENTAL_REPORT.open("w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(
            f,
            fieldnames=["檔案路徑", "原狀態", "新狀態", "處理動作", "輸出檔", "說明"],
        )
        writer.writeheader()
        writer.writerows(supplemental_rows)

    removed = update_checkpoint_remove(rewritten_outputs)
    print(f"Processed rows: {len(supplemental_rows)}")
    print(f"Rewritten outputs: {len(rewritten_outputs)}")
    print(f"Removed checkpoint entries: {removed}")
    print(f"Supplemental report: {SUPPLEMENTAL_REPORT}")


if __name__ == "__main__":
    main()
